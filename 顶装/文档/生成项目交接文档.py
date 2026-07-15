from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "云帆瑞达顶装Lua项目交接文档_v3.1.md"
OUTPUT = BASE / "云帆瑞达顶装Lua项目交接文档_v3.1.docx"

NAVY = "1F3864"
BLUE = "2E75B6"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
BODY = "333333"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.name = "Microsoft YaHei"
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    paragraph.add_run(" 页")


def add_toc(paragraph):
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run._r.append(instr)
    sep_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(separate)
    placeholder = paragraph.add_run("打开 Word 后右键更新目录")
    placeholder.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_inline(paragraph, text, bold=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Cascadia Mono"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        else:
            run = paragraph.add_run(part)
            run.bold = bold


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, outline in (
        ("Heading 1", 18, NAVY, 0),
        ("Heading 2", 14, BLUE, 1),
        ("Heading 3", 12, NAVY, 2),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(18 if outline == 0 else 12)
        style.paragraph_format.space_after = Pt(8 if outline == 0 else 5)
        p_pr = style.element.get_or_add_pPr()
        outline_el = p_pr.find(qn("w:outlineLvl"))
        if outline_el is None:
            outline_el = OxmlElement("w:outlineLvl")
            p_pr.append(outline_el)
        outline_el.set(qn("w:val"), str(outline))

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Cascadia Mono"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    code.paragraph_format.left_indent = Cm(0.5)
    code.paragraph_format.right_indent = Cm(0.5)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)


def add_cover(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("云帆瑞达 60G 毫米波雷达")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("顶装 Lua 项目交接文档")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("人数统计 · 三人轨迹 · 站坐躺姿态 · 上床下桌区域")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph("")
    meta = [line.replace("  ", "").strip() for line in lines if line.strip()]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        add_inline(p, line)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("交接状态：代码与离线验证已完成，现场参数待标定")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()

    toc_title = doc.add_paragraph(style="Heading 1")
    toc_title.add_run("目录")
    toc = doc.add_paragraph()
    add_toc(toc)
    doc.add_page_break()


def add_table(doc, rows):
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, bold=(r_idx == 0))
            for run in p.runs:
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header.paragraphs[0]
    header.text = "云帆瑞达顶装 Lua 项目交接文档 · Radar_Top_Posture_3.1"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = "云帆瑞达60G毫米波雷达顶装Lua项目交接文档"
    doc.core_properties.subject = "人数统计、三人轨迹和站坐躺姿态模块交接"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "毫米波雷达,Lua,人数统计,轨迹跟踪,姿态识别,项目交接"

    first_h2 = next(i for i, line in enumerate(lines) if line.startswith("## "))
    cover_meta = [line for line in lines[1:first_h2] if not line.startswith("#")]
    add_cover(doc, cover_meta)

    i = first_h2
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                p.paragraph_format.keep_together = True
                set_shd = OxmlElement("w:shd")
                set_shd.set(qn("w:fill"), "F5F7FA")
                p._p.get_or_add_pPr().append(set_shd)
                p.add_run("\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("| ") and i + 1 < len(lines) and re.match(r"^\|[-: |]+\|$", lines[i + 1]):
            rows = []
            rows.append([c.strip() for c in line.strip("|").split("|")])
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            if p._p.getprevious() is not None:
                p.paragraph_format.page_break_before = True
            add_inline(p, line[3:])
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:])
        elif line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[5:])
        elif re.match(r"^- \[[ xX]\] ", line):
            checked = line[3].lower() == "x"
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, ("☑ " if checked else "☐ ") + line[6:])
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, line)
        i += 1

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.save(OUTPUT)
    print(f"CREATED: {OUTPUT}")
    print(f"PARAGRAPHS: {len(doc.paragraphs)}")
    print(f"TABLES: {len(doc.tables)}")


if __name__ == "__main__":
    build()
